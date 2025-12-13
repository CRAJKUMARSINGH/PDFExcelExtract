"""
Examine File Content
Checks what's actually in the converted files
"""

import os
from pathlib import Path

def examine_word_file(file_path):
    """Examine content of a Word file"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        print(f"Examining: {Path(file_path).name}")
        print("-" * 40)
        
        # Show document structure
        print(f"Document has {len(doc.paragraphs)} paragraphs")
        print(f"Document has {len(doc.tables)} tables")
        
        # Show first few paragraphs
        print("\nFirst 5 paragraphs:")
        for i, para in enumerate(doc.paragraphs[:5]):
            text = para.text.strip()
            if text:
                print(f"  {i+1}: {text[:60]}{'...' if len(text) > 60 else ''}")
        
        # Show first table if exists
        if doc.tables:
            print(f"\nFirst table ({len(doc.tables[0].rows)} rows x {len(doc.tables[0].columns)} columns):")
            table = doc.tables[0]
            for i, row in enumerate(table.rows[:3]):  # First 3 rows
                row_text = " | ".join([cell.text.strip()[:20] for cell in row.cells])
                print(f"  Row {i+1}: {row_text}")
        else:
            print("\nNo tables found in document")
            
        print()
        return True
    except Exception as e:
        print(f"Error examining {file_path}: {e}")
        return False

def examine_excel_file(file_path):
    """Examine content of an Excel file"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path)
        ws = wb.active
        
        print(f"Examining: {Path(file_path).name}")
        print("-" * 40)
        
        # Show sheet info
        print(f"Sheet name: {ws.title}")
        print(f"Dimensions: {ws.max_row} rows x {ws.max_column} columns")
        
        # Show first few rows
        print("\nFirst 5 rows:")
        for row_num in range(1, min(6, ws.max_row + 1)):
            row_data = []
            for col_num in range(1, min(6, ws.max_column + 1)):  # First 5 columns
                cell = ws.cell(row=row_num, column=col_num)
                value = str(cell.value)[:30] if cell.value is not None else ""
                row_data.append(value)
            print(f"  Row {row_num}: {' | '.join(row_data)}")
        
        # Check for actual data vs headers
        data_rows = 0
        for row_num in range(1, min(20, ws.max_row + 1)):
            empty_cells = 0
            for col_num in range(1, min(10, ws.max_column + 1)):
                cell = ws.cell(row=row_num, column=col_num)
                if cell.value is None or str(cell.value).strip() == "":
                    empty_cells += 1
            if empty_cells < 5:  # If most cells in row have data
                data_rows += 1
        
        print(f"\nData density: {data_rows}/20 rows with significant content")
        print()
        
        wb.close()
        return True
    except Exception as e:
        print(f"Error examining {file_path}: {e}")
        return False

def main():
    """Main function"""
    print("FILE CONTENT EXAMINATION")
    print("=" * 25)
    
    # Examine one of the problematic files
    word_file = "output/ACE Scanner_2025_11_19(2).docx"
    excel_file = "output/ACE Scanner_2025_11_19(2).xlsx"
    
    if os.path.exists(word_file):
        examine_word_file(word_file)
    
    if os.path.exists(excel_file):
        examine_excel_file(excel_file)
    
    # Also examine one of the better files
    print("\n" + "="*50)
    print("COMPARISON WITH IMPROVED VERSION")
    print("="*50)
    
    improved_word = "improved_output/ACE Scanner_2025_11_19(2)_FORMATTED.docx"
    improved_excel = "improved_output/ACE Scanner_2025_11_19(2)_TABLES.xlsx"
    
    if os.path.exists(improved_word):
        examine_word_file(improved_word)
    
    if os.path.exists(improved_excel):
        examine_excel_file(improved_excel)

if __name__ == "__main__":
    main()