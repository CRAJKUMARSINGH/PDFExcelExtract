"""
Check OCR Results
Examines the quality of OCR-converted files
"""

import os
from pathlib import Path

def examine_word_file(file_path):
    """Examine content of a Word file"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        print(f"Examining: {Path(file_path).name}")
        print("-" * 40)
        
        # Show document structure
        print(f"Document has {len(doc.paragraphs)} paragraphs")
        print(f"Document has {len(doc.tables)} tables")
        
        # Show first few paragraphs
        print("\nFirst 5 paragraphs:")
        paragraphs_shown = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text and paragraphs_shown < 5:
                print(f"  {paragraphs_shown+1}: {text[:60]}{'...' if len(text) > 60 else ''}")
                paragraphs_shown += 1
        
        # Show last few paragraphs if document is long
        if len(doc.paragraphs) > 10:
            print("\nLast 3 paragraphs:")
            paragraphs_shown = 0
            for i in range(len(doc.paragraphs)-1, -1, -1):
                text = doc.paragraphs[i].text.strip()
                if text and paragraphs_shown < 3:
                    print(f"  {paragraphs_shown+1}: {text[:60]}{'...' if len(text) > 60 else ''}")
                    paragraphs_shown += 1
        
        print()
        return True
    except Exception as e:
        print(f"Error examining {file_path}: {e}")
        return False

def examine_excel_file(file_path):
    """Examine content of an Excel file"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path)
        ws = wb.active
        
        print(f"Examining: {Path(file_path).name}")
        print("-" * 40)
        
        # Show sheet info
        print(f"Sheet name: {ws.title}")
        print(f"Dimensions: {ws.max_row} rows x {ws.max_column} columns")
        
        # Check for actual data
        data_rows = 0
        sample_rows = []
        for row_num in range(1, min(20, ws.max_row + 1)):
            row_data = []
            empty_cells = 0
            for col_num in range(1, min(10, ws.max_column + 1)):
                cell = ws.cell(row=row_num, column=col_num)
                value = str(cell.value)[:30] if cell.value is not None else ""
                row_data.append(value)
                if not value:
                    empty_cells += 1
            
            # If row has significant content
            if empty_cells < 7:
                data_rows += 1
                if len(sample_rows) < 5:
                    sample_rows.append((row_num, row_data))
        
        print(f"Data density: {data_rows}/20 rows with significant content")
        
        # Show sample rows with data
        if sample_rows:
            print("\nSample data rows:")
            for row_num, row_data in sample_rows[:3]:
                print(f"  Row {row_num}: {' | '.join(row_data)}")
        
        print()
        
        wb.close()
        return True
    except Exception as e:
        print(f"Error examining {file_path}: {e}")
        return False

def main():
    """Main function"""
    print("OCR RESULTS EXAMINATION")
    print("=" * 25)
    
    # Check the OCR output directory
    output_dir = "ocr_output_final"
    
    if not os.path.exists(output_dir):
        print(f"Directory {output_dir} not found")
        return
    
    # Find some files to examine
    word_files = []
    excel_files = []
    
    for file in os.listdir(output_dir):
        if file.endswith('_OCR.docx'):
            word_files.append(os.path.join(output_dir, file))
        elif file.endswith('_OCR.xlsx'):
            excel_files.append(os.path.join(output_dir, file))
    
    print(f"Found {len(word_files)} Word files and {len(excel_files)} Excel files")
    print()
    
    # Examine a few representative files
    if word_files:
        print("WORD DOCUMENT ANALYSIS:")
        print("=" * 25)
        # Examine a couple of files
        for word_file in word_files[:2]:
            examine_word_file(word_file)
    
    if excel_files:
        print("EXCEL DOCUMENT ANALYSIS:")
        print("=" * 25)
        # Examine a couple of files
        for excel_file in excel_files[:2]:
            examine_excel_file(excel_file)

if __name__ == "__main__":
    main()