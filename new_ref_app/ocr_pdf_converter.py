"""
OCR PDF Converter
Converts image-based PDFs to Word and Excel using Optical Character Recognition
"""

import os
import sys
from pathlib import Path

def install_required_packages():
    """Install required packages for OCR"""
    try:
        import pytesseract
        from PIL import Image
        import pdf2image
        import pandas as pd
        from docx import Document
        import openpyxl
        print("Required OCR packages already installed.")
        return True
    except ImportError:
        print("Installing OCR packages...")
        try:
            import subprocess
            import sys
            subprocess.check_call([sys.executable, "-m", "pip", "install", "pytesseract", "pdf2image", "pandas", "python-docx", "openpyxl", "Pillow"])
            print("OCR packages installed successfully.")
            return True
        except Exception as e:
            print(f"Failed to install OCR packages: {e}")
            return False

def check_tesseract():
    """Check if Tesseract OCR is installed"""
    try:
        import pytesseract
        # Try to get tesseract version
        version = pytesseract.get_tesseract_version()
        print(f"Tesseract OCR version: {version}")
        return True
    except Exception as e:
        print(f"Tesseract OCR not found: {e}")
        print("Please install Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki")
        return False

def get_pdf_files(directory):
    """Get all PDF files in the directory"""
    pdf_files = []
    for file in os.listdir(directory):
        if file.lower().endswith('.pdf'):
            pdf_files.append(file)
    return pdf_files

def ocr_pdf_to_text(pdf_path):
    """Convert PDF to text using OCR"""
    try:
        import pytesseract
        from pdf2image import convert_from_path
        import tempfile
        
        print("  Converting PDF to images for OCR...")
        # Convert PDF pages to images (without specifying poppler_path)
        try:
            pages = convert_from_path(pdf_path, dpi=300, grayscale=True)
        except Exception as e:
            print(f"  Poppler error, trying alternative method: {e}")
            # Try with lower DPI
            pages = convert_from_path(pdf_path, dpi=200, grayscale=True)
        
        print(f"  Processing {len(pages)} pages with OCR...")
        text_content = ""
        
        for i, page in enumerate(pages):
            # Apply OCR to each page
            page_text = pytesseract.image_to_string(page)
            text_content += f"\n\n--- PAGE {i+1} ---\n\n"
            text_content += page_text
        
        return text_content if text_content.strip() else "[OCR completed but no text found]"
    except Exception as e:
        print(f"  OCR Error: {e}")
        return f"[OCR Error: {str(e)}]"

def create_word_with_ocr_content(content, output_path, original_filename):
    """Create Word document with OCR content"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Ensure the output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'OCR CONVERTED: {original_filename}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('')
        doc.add_paragraph(f'Converted using Optical Character Recognition (OCR)')
        doc.add_paragraph('')
        doc.add_page_break()
        
        # Process content
        if "[OCR Error]" in content or "[OCR completed but no text found]" in content:
            doc.add_paragraph(content)
        else:
            # Split content into sections/pages
            pages = content.split('--- PAGE')
            for i, page_content in enumerate(pages):
                if i == 0 and not page_content.strip():  # Skip first empty section
                    continue
                    
                if i > 0:
                    doc.add_heading(f'Page {i}', level=1)
                
                # Add paragraphs
                paragraphs = page_content.split('\n\n')
                for para in paragraphs:
                    if para.strip() and "PAGE" not in para:  # Skip page markers
                        doc.add_paragraph(para.strip())
                
                # Add page break (except for last page)
                if i < len(pages) - 1 and i > 0:
                    doc.add_page_break()
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating Word document: {e}")
        # Try creating a simple text file as fallback
        try:
            with open(output_path.replace('.docx', '_fallback.txt'), 'w', encoding='utf-8') as f:
                f.write(f"OCR CONVERTED: {original_filename}\n\n")
                f.write(content)
            print(f"  Created fallback text file instead")
            return True
        except Exception as e2:
            print(f"  Fallback also failed: {e2}")
            return False

def create_excel_with_ocr_content(content, output_path, original_filename):
    """Create Excel document with OCR content"""
    try:
        from openpyxl import Workbook
        import re
        
        # Ensure the output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        wb = Workbook()
        ws = wb.active
        ws.title = "OCR Content"
        
        # Add header
        ws['A1'] = f'OCR CONVERTED: {original_filename}'
        ws['A2'] = 'Converted using Optical Character Recognition'
        ws['A3'] = ''
        
        current_row = 4
        
        if "[OCR Error]" in content or "[OCR completed but no text found]" in content:
            ws[f'A{current_row}'] = content
        else:
            # Try to identify tabular data
            lines = content.split('\n')
            for line in lines:
                if line.strip() and not line.startswith('--- PAGE'):
                    # Look for tabular patterns
                    if '\t' in line or re.search(r'\s{2,}', line):
                        # Likely tabular data
                        if '\t' in line:
                            cells = line.split('\t')
                        else:
                            cells = re.split(r'\s{2,}', line)  # Split by 2+ spaces
                        
                        # Write to consecutive columns
                        for col_num, cell_value in enumerate(cells, 1):
                            ws.cell(row=current_row, column=col_num, value=cell_value.strip())
                    else:
                        # Regular text
                        ws.cell(row=current_row, column=1, value=line.strip())
                    
                    current_row += 1
        
        # Auto-adjust column widths
        for col in range(1, 11):  # Adjust first 10 columns
            col_letter = chr(64 + col)
            ws.column_dimensions[col_letter].width = 25
        
        wb.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating Excel document: {e}")
        # Try creating a simple text file as fallback
        try:
            with open(output_path.replace('.xlsx', '_fallback.txt'), 'w', encoding='utf-8') as f:
                f.write(f"OCR CONVERTED: {original_filename}\n\n")
                f.write(content)
            print(f"  Created fallback text file instead")
            return True
        except Exception as e2:
            print(f"  Fallback also failed: {e2}")
            return False

def convert_pdfs_with_ocr(input_dir, output_dir):
    """Convert all PDF files using OCR"""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Get list of PDF files
    pdf_files = get_pdf_files(input_dir)
    
    if not pdf_files:
        print("No PDF files found in the directory.")
        return False
    
    print(f"Found {len(pdf_files)} PDF files to convert with OCR:")
    for pdf_file in pdf_files:
        print(f"  - {pdf_file}")
    
    # Process each PDF file
    successful_conversions = 0
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}] Processing: {pdf_file}")
        
        pdf_path = os.path.join(input_dir, pdf_file)
        base_name = Path(pdf_file).stem
        
        # Apply OCR to extract text
        print("  Applying OCR to extract text...")
        content = ocr_pdf_to_text(pdf_path)
        
        # Create Word document
        word_output = os.path.join(output_dir, f"{base_name}_OCR.docx")
        print("  Creating OCR Word document...")
        if create_word_with_ocr_content(content, word_output, pdf_file):
            print(f"  ✓ Created OCR Word document")
        else:
            print(f"  ✗ Failed to create OCR Word document")
            continue
        
        # Create Excel document
        excel_output = os.path.join(output_dir, f"{base_name}_OCR.xlsx")
        print("  Creating OCR Excel document...")
        if create_excel_with_ocr_content(content, excel_output, pdf_file):
            print(f"  ✓ Created OCR Excel document")
        else:
            print(f"  ✗ Failed to create OCR Excel document")
            continue
        
        successful_conversions += 1
    
    print(f"\n🎉 OCR conversion completed!")
    print(f"Successfully processed {successful_conversions} out of {len(pdf_files)} files.")
    print(f"Output saved to: {output_dir}")
    return True

def main():
    """Main function"""
    print("OCR PDF to Word/Excel Converter")
    print("=" * 32)
    
    # Install required packages
    if not install_required_packages():
        print("Cannot proceed without required OCR packages.")
        return False
    
    # Check if Tesseract is installed
    if not check_tesseract():
        print("\nTesseract OCR is required for this converter.")
        print("Please install it from: https://github.com/UB-Mannheim/tesseract/wiki")
        print("After installation, add it to your system PATH.")
        return False
    
    # Import packages after installation
    try:
        import pytesseract
        from PIL import Image
        import pdf2image
        import pandas as pd
        from docx import Document
        import openpyxl
    except ImportError as e:
        print(f"Failed to import OCR packages: {e}")
        return False
    
    # Set directories
    input_dir = os.getcwd()  # Current directory
    output_dir = os.path.join(input_dir, "ocr_output_final")
    
    print(f"Input directory: {input_dir}")
    print(f"Output directory: {output_dir}")
    print()
    
    try:
        success = convert_pdfs_with_ocr(input_dir, output_dir)
        return success
    except Exception as e:
        print(f"Error during OCR conversion: {str(e)}")
        return False

if __name__ == "__main__":
    success = main()
    if success:
        print("\n✅ OCR conversion completed successfully!")
        print("📁 Check the 'ocr_output_final' folder for your documents.")
    else:
        print("\n❌ OCR conversion process encountered errors.")
    sys.exit(0 if success else 1)