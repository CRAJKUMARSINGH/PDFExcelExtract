"""
Advanced PDF Converter with Table Detection
Converts PDF files to Excel with proper table formatting
"""

import os
import sys
from pathlib import Path
import re

def install_required_packages():
    """Install required packages if not already installed"""
    try:
        import pdfplumber
        import pandas as pd
        from docx import Document
        import openpyxl
        print("Required packages already installed.")
        return True
    except ImportError:
        print("Installing required packages...")
        try:
            import subprocess
            import sys
            subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber", "pandas", "python-docx", "openpyxl"])
            print("Packages installed successfully.")
            return True
        except Exception as e:
            print(f"Failed to install packages: {e}")
            return False

def get_pdf_files(directory):
    """Get all PDF files in the directory"""
    pdf_files = []
    for file in os.listdir(directory):
        if file.lower().endswith('.pdf'):
            pdf_files.append(file)
    return pdf_files

def convert_pdf_to_excel_with_tables(pdf_path, output_path):
    """Convert PDF to Excel with proper table formatting"""
    try:
        import pdfplumber
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.utils.dataframe import dataframe_to_rows
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "PDF Tables"
        
        # Add header
        filename = Path(pdf_path).name
        ws['A1'] = f'TABLE DATA FROM: {filename}'
        ws['A2'] = 'This worksheet contains properly formatted tables extracted from the PDF'
        ws['A3'] = ''
        
        current_row = 4
        
        # Try to extract tables
        with pdfplumber.open(pdf_path) as pdf:
            table_count = 0
            for page_num, page in enumerate(pdf.pages, 1):
                # Try to extract tables from page
                tables = page.extract_tables()
                
                if tables:
                    # Process tables
                    for table_idx, table in enumerate(tables):
                        table_count += 1
                        # Add table header
                        ws[f'A{current_row}'] = f'TABLE {table_count} (Page {page_num})'
                        current_row += 1
                        
                        if table and len(table) > 0:
                            # Convert table to DataFrame
                            # Handle cases where table might have inconsistent row lengths
                            max_cols = max(len(row) for row in table) if table else 0
                            # Pad rows to have the same number of columns
                            padded_table = []
                            for row in table:
                                padded_row = row + [None] * (max_cols - len(row)) if len(row) < max_cols else row
                                padded_table.append(padded_row)
                            
                            if padded_table:
                                try:
                                    df = pd.DataFrame(padded_table[1:], columns=padded_table[0]) if len(padded_table) > 1 else pd.DataFrame([padded_table[0]])
                                    
                                    # Add table to worksheet
                                    for r_idx, r in enumerate(dataframe_to_rows(df, index=False, header=True)):
                                        for col_num, value in enumerate(r, 1):
                                            # Clean up None values
                                            cell_value = "" if value is None else str(value)
                                            ws.cell(row=current_row, column=col_num, value=cell_value)
                                        current_row += 1
                                except Exception as e:
                                    # Fallback: add raw table data
                                    for row_data in padded_table:
                                        for col_num, cell_data in enumerate(row_data, 1):
                                            cell_value = "" if cell_data is None else str(cell_data)
                                            ws.cell(row=current_row, column=col_num, value=cell_value)
                                        current_row += 1
                        
                        # Add spacing
                        current_row += 2
                else:
                    # Extract text if no tables found
                    text = page.extract_text()
                    if text:
                        # Try to identify potential tabular data in text
                        lines = text.split('\n')
                        # Look for lines with multiple tab-separated or space-separated values
                        tabular_lines = []
                        for line in lines:
                            # Count tabs or multiple spaces
                            if line.count('\t') > 1 or line.count('  ') > 2:
                                tabular_lines.append(line)
                        
                        if tabular_lines:
                            ws[f'A{current_row}'] = f'POTENTIAL TABULAR DATA (Page {page_num})'
                            current_row += 1
                            
                            # Try to parse as table
                            for line in tabular_lines:
                                # Split by tabs or multiple spaces
                                if '\t' in line:
                                    cells = line.split('\t')
                                else:
                                    cells = re.split(r'  +', line)  # Split by 2 or more spaces
                                
                                for col_num, cell_value in enumerate(cells, 1):
                                    ws.cell(row=current_row, column=col_num, value=cell_value.strip())
                                current_row += 1
                            
                            current_row += 1
        
        # Auto-adjust column widths
        for col in range(1, 11):  # Adjust first 10 columns
            col_letter = chr(64 + col)  # A, B, C, ...
            ws.column_dimensions[col_letter].width = 25
        
        # Save workbook
        wb.save(output_path)
        print(f"  Created Excel with {table_count} tables detected")
        return True
    except Exception as e:
        print(f"Error converting PDF to Excel: {e}")
        # Create a simple error file
        with open(output_path.replace('.xlsx', '_error.txt'), 'w') as f:
            f.write(f"Error converting {Path(pdf_path).name} to Excel:\n{str(e)}")
        return False

def convert_pdf_to_word_formatted(pdf_path, output_path):
    """Convert PDF to formatted Word document"""
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        filename = Path(pdf_path).name
        heading = doc.add_heading(f'TABLE DATA FROM: {filename}', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('')
        doc.add_paragraph('This document contains properly formatted tables extracted from the PDF.')
        doc.add_paragraph('')
        
        # Process PDF
        with pdfplumber.open(pdf_path) as pdf:
            table_count = 0
            for page_num, page in enumerate(pdf.pages, 1):
                # Try to extract tables
                tables = page.extract_tables()
                
                if tables:
                    # Process tables
                    for table_idx, table_data in enumerate(tables):
                        if table_data and len(table_data) > 0:
                            table_count += 1
                            # Add table header
                            doc.add_heading(f'Table {table_count} (Page {page_num})', level=1)
                            
                            # Create table in Word
                            max_cols = max(len(row) for row in table_data) if table_data else 1
                            # Pad rows to have the same number of columns
                            padded_table = []
                            for row in table_data:
                                padded_row = row + [None] * (max_cols - len(row)) if len(row) < max_cols else row
                                padded_table.append(padded_row)
                            
                            if padded_table:
                                table = doc.add_table(rows=len(padded_table), cols=max_cols)
                                table.style = 'Table Grid'
                                
                                # Fill table
                                for i, row_data in enumerate(padded_table):
                                    row = table.rows[i]
                                    for j, cell_data in enumerate(row_data):
                                        if j < len(row.cells):
                                            cell_text = "" if cell_data is None else str(cell_data)
                                            row.cells[j].text = cell_text
                
                # Extract text
                text = page.extract_text()
                if text:
                    # Add page marker
                    doc.add_heading(f'Page {page_num} Text Content', level=2)
                    
                    # Add paragraphs
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                
                # Add page break (except for last page)
                if page_num < len(pdf.pages):
                    doc.add_page_break()
        
        doc.save(output_path)
        print(f"  Created Word document with {table_count} tables")
        return True
    except Exception as e:
        print(f"Error converting PDF to Word: {e}")
        # Create a simple error file
        with open(output_path.replace('.docx', '_error.txt'), 'w') as f:
            f.write(f"Error converting {Path(pdf_path).name} to Word:\n{str(e)}")
        return False

def convert_pdfs_advanced(input_dir, output_dir):
    """Convert all PDF files with advanced formatting preservation"""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Get list of PDF files
    pdf_files = get_pdf_files(input_dir)
    
    if not pdf_files:
        print("No PDF files found in the directory.")
        return False
    
    print(f"Found {len(pdf_files)} PDF files to convert:")
    for pdf_file in pdf_files:
        print(f"  - {pdf_file}")
    
    # Process each PDF file
    successful_conversions = 0
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}] Processing: {pdf_file}")
        
        pdf_path = os.path.join(input_dir, pdf_file)
        base_name = Path(pdf_file).stem
        
        # Convert to formatted Word
        word_output = os.path.join(output_dir, f"{base_name}_FORMATTED.docx")
        print("  Creating formatted Word document...")
        if convert_pdf_to_word_formatted(pdf_path, word_output):
            print(f"  ✓ Created formatted Word document")
        else:
            print(f"  ✗ Failed to create formatted Word document")
        
        # Convert to structured Excel
        excel_output = os.path.join(output_dir, f"{base_name}_TABLES.xlsx")
        print("  Creating structured Excel document...")
        if convert_pdf_to_excel_with_tables(pdf_path, excel_output):
            print(f"  ✓ Created structured Excel document")
        else:
            print(f"  ✗ Failed to create structured Excel document")
        
        successful_conversions += 1
    
    print(f"\n🎉 Advanced conversion completed!")
    print(f"Processed {successful_conversions} out of {len(pdf_files)} files.")
    print(f"Output saved to: {output_dir}")
    return True

def main():
    """Main function"""
    print("Advanced PDF to Word/Excel Converter")
    print("=" * 38)
    
    # Install required packages
    if not install_required_packages():
        print("Cannot proceed without required packages.")
        return False
    
    # Import packages after installation
    try:
        import pdfplumber
        import pandas as pd
        from docx import Document
        import openpyxl
    except ImportError as e:
        print(f"Failed to import packages: {e}")
        return False
    
    # Set directories
    input_dir = os.getcwd()  # Current directory
    output_dir = os.path.join(input_dir, "improved_output")
    
    print(f"Input directory: {input_dir}")
    print(f"Output directory: {output_dir}")
    print()
    
    try:
        success = convert_pdfs_advanced(input_dir, output_dir)
        return success
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        return False

if __name__ == "__main__":
    success = main()
    if success:
        print("\n✅ Advanced conversion completed successfully!")
        print("📁 Check the 'improved_output' folder for your documents.")
    else:
        print("\n❌ Advanced conversion process encountered errors.")
    sys.exit(0 if success else 1)