"""
PDF to Word and Excel Converter
Converts PDF files to both Word (.docx) and Excel (.xlsx) formats
Saves output in a dedicated 'output' folder
"""

import os
import sys
from pathlib import Path
import traceback

def install_required_packages():
    """Install required packages if not already installed"""
    try:
        import PyPDF2
        from docx import Document
        import openpyxl
        print("Required packages already installed.")
        return True
    except ImportError:
        print("Installing required packages...")
        try:
            import subprocess
            import sys
            subprocess.check_call([sys.executable, "-m", "pip", "install", "PyPDF2", "python-docx", "openpyxl"])
            print("Packages installed successfully.")
            return True
        except Exception as e:
            print(f"Failed to install packages: {e}")
            return False

def get_pdf_files(directory):
    """Get all PDF files in the directory"""
    pdf_files = []
    for file in os.listdir(directory):
        if file.lower().endswith('.pdf'):
            pdf_files.append(file)
    return pdf_files

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file"""
    try:
        import PyPDF2
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                except:
                    continue
        return text if text.strip() else "[Could not extract text - may be image-only PDF]"
    except Exception as e:
        return f"[Error extracting text: {str(e)}]"

def create_word_document(content, output_path, original_filename):
    """Create a Word document from content"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Add title
        doc.add_heading('Converted PDF Document', 0)
        
        # Add file information
        doc.add_paragraph(f'Original file: {original_filename}')
        doc.add_paragraph('')
        doc.add_page_break()
        
        # Add content
        if content.startswith("[") and content.endswith("]"):
            # Error or warning message
            doc.add_paragraph(content)
        else:
            # Normal content - split into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return False

def create_excel_document(content, output_path, original_filename):
    """Create an Excel document from content"""
    try:
        from openpyxl import Workbook
        
        wb = Workbook()
        ws = wb.active
        ws.title = "PDF Content"
        
        # Add header
        ws['A1'] = 'Converted PDF Content'
        ws['A2'] = f'Original file: {original_filename}'
        ws['A3'] = ''
        
        # Add content
        if content.startswith("[") and content.endswith("]"):
            # Error or warning message
            ws['A4'] = content
        else:
            # Normal content
            lines = content.split('\n')
            row = 4
            for line in lines:
                if line.strip():
                    ws[f'A{row}'] = line.strip()
                    row += 1
        
        # Auto-adjust column width
        ws.column_dimensions['A'].width = 50
        
        wb.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating Excel document: {e}")
        return False

def convert_pdfs(input_dir, output_dir):
    """Convert all PDF files in input directory to Word and Excel"""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Get list of PDF files
    pdf_files = get_pdf_files(input_dir)
    
    if not pdf_files:
        print("No PDF files found in the directory.")
        return False
    
    print(f"Found {len(pdf_files)} PDF files to convert:")
    for pdf_file in pdf_files:
        print(f"  - {pdf_file}")
    
    # Process each PDF file
    successful_conversions = 0
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}] Processing: {pdf_file}")
        
        pdf_path = os.path.join(input_dir, pdf_file)
        base_name = Path(pdf_file).stem
        
        # Extract text from PDF
        print("  Extracting text from PDF...")
        content = extract_text_from_pdf(pdf_path)
        
        # Convert to Word
        word_output = os.path.join(output_dir, f"{base_name}.docx")
        print("  Creating Word document...")
        if create_word_document(content, word_output, pdf_file):
            print(f"  ✓ Created Word document: {os.path.basename(word_output)}")
        else:
            print(f"  ✗ Failed to create Word document")
            continue
        
        # Convert to Excel
        excel_output = os.path.join(output_dir, f"{base_name}.xlsx")
        print("  Creating Excel document...")
        if create_excel_document(content, excel_output, pdf_file):
            print(f"  ✓ Created Excel document: {os.path.basename(excel_output)}")
        else:
            print(f"  ✗ Failed to create Excel document")
            continue
        
        successful_conversions += 1
    
    print(f"\n🎉 Conversion completed!")
    print(f"Successfully converted {successful_conversions} out of {len(pdf_files)} files.")
    print(f"Output saved to: {output_dir}")
    return True

def main():
    """Main function"""
    print("PDF to Word and Excel Converter")
    print("=" * 32)
    
    # Install required packages
    if not install_required_packages():
        print("Cannot proceed without required packages.")
        return False
    
    # Import packages after installation
    try:
        import PyPDF2
        from docx import Document
        import openpyxl
    except ImportError as e:
        print(f"Failed to import packages: {e}")
        return False
    
    # Set directories
    input_dir = os.getcwd()  # Current directory
    output_dir = os.path.join(input_dir, "output")
    
    print(f"Input directory: {input_dir}")
    print(f"Output directory: {output_dir}")
    print()
    
    try:
        success = convert_pdfs(input_dir, output_dir)
        return success
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = main()
    if success:
        print("\n✅ All files converted successfully!")
        print("📁 Check the 'output' folder for your converted files.")
    else:
        print("\n❌ Conversion process encountered errors.")
    sys.exit(0 if success else 1)