"""
Batch PDF Converter
Converts all PDF files in a directory to both Word and Excel formats
Uses proper libraries for text extraction and document creation
"""

import os
import sys
from pathlib import Path
import traceback

def install_and_import_packages():
    """Install and import required packages"""
    try:
        import PyPDF2
        import pdfplumber
        from docx import Document
        import openpyxl
        from openpyxl import Workbook
        print("All required packages are already installed.")
        return True
    except ImportError as e:
        print(f"Missing package: {e}")
        print("Installing required packages...")
        try:
            import subprocess
            import sys
            
            # Try pip first
            subprocess.check_call([sys.executable, "-m", "pip", "install", "PyPDF2", "pdfplumber", "python-docx", "openpyxl"])
            print("Packages installed successfully.")
            return True
        except Exception as install_error:
            print(f"Failed to install packages: {install_error}")
            print("Please run 'pip install PyPDF2 pdfplumber python-docx openpyxl' manually.")
            return False

def get_pdf_files(directory):
    """Get all PDF files in the directory"""
    pdf_files = []
    for file in os.listdir(directory):
        if file.lower().endswith('.pdf'):
            pdf_files.append(file)
    return pdf_files

def convert_pdf_to_text(pdf_path):
    """
    Extract text from PDF file using multiple methods for better accuracy
    """
    try:
        text = ""
        
        # Method 1: Try pdfplumber (better for tables and layouts)
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            if text.strip():
                print("  Used pdfplumber for text extraction")
                return text
        except Exception as e:
            print(f"  pdfplumber failed: {e}")
        
        # Method 2: Try PyPDF2 as fallback
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            if text.strip():
                print("  Used PyPDF2 for text extraction")
                return text
        except Exception as e:
            print(f"  PyPDF2 failed: {e}")
        
        # If both methods fail
        if not text.strip():
            return f"[Could not extract text from {os.path.basename(pdf_path)} - may be image-only PDF requiring OCR]"
        
        return text
    except Exception as e:
        print(f"Error extracting text from {pdf_path}: {str(e)}")
        return f"[Error extracting text: {str(e)}]"

def create_word_document(text, output_path):
    """
    Create a Word document from text with basic formatting
    """
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Add title
        doc.add_heading('Converted PDF Document', 0)
        
        # Add metadata
        doc.add_paragraph(f'Original file: {Path(output_path).stem}.pdf')
        doc.add_paragraph('Converted using PDF Batch Converter')
        doc.add_page_break()
        
        # Add content
        # Split text into paragraphs (assuming double newlines separate paragraphs)
        paragraphs = text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save document
        doc.save(output_path)
        print(f"  Created Word document: {os.path.basename(output_path)}")
        return True
    except Exception as e:
        print(f"Error creating Word document {output_path}: {str(e)}")
        return False

def create_excel_spreadsheet(text, output_path):
    """
    Create an Excel spreadsheet from text
    Attempts to preserve table-like structures
    """
    try:
        from openpyxl import Workbook
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Converted PDF Content"
        
        # Add header information
        ws['A1'] = 'Converted PDF Content'
        ws['A2'] = f'Original file: {Path(output_path).stem}.pdf'
        ws['A3'] = 'Converted using PDF Batch Converter'
        ws['A4'] = ''
        
        # Process text content
        lines = text.split('\n')
        row_index = 5
        
        for line in lines:
            if line.strip():
                # Try to detect tabular data (lines with tabs or multiple spaces)
                if '\t' in line or '  ' in line:
                    # Split by tabs or multiple spaces
                    if '\t' in line:
                        cells = line.split('\t')
                    else:
                        cells = line.split('  ')
                    
                    # Write to consecutive columns
                    for col_index, cell_content in enumerate(cells, 1):
                        ws.cell(row=row_index, column=col_index, value=cell_content.strip())
                else:
                    # Regular text line
                    ws.cell(row=row_index, column=1, value=line.strip())
                
                row_index += 1
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Save workbook
        wb.save(output_path)
        print(f"  Created Excel spreadsheet: {os.path.basename(output_path)}")
        return True
    except Exception as e:
        print(f"Error creating Excel spreadsheet {output_path}: {str(e)}")
        return False

def process_pdf_files(input_dir, output_dir):
    """Process all PDF files in the input directory"""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Get list of PDF files
    pdf_files = get_pdf_files(input_dir)
    
    if not pdf_files:
        print("No PDF files found in the directory.")
        return False
    
    print(f"Found {len(pdf_files)} PDF files to convert:")
    for pdf_file in pdf_files:
        print(f"  - {pdf_file}")
    
    # Process each PDF file
    successful_conversions = 0
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}] Processing: {pdf_file}")
        
        pdf_path = os.path.join(input_dir, pdf_file)
        base_name = Path(pdf_file).stem
        
        # Extract text from PDF
        print("  Extracting text from PDF...")
        text_content = convert_pdf_to_text(pdf_path)
        
        if not text_content:
            print(f"  Failed to extract content from {pdf_file}")
            continue
        
        # Convert to Word
        word_output = os.path.join(output_dir, f"{base_name}.docx")
        print("  Creating Word document...")
        if create_word_document(text_content, word_output):
            print(f"  ✓ Created Word document")
        else:
            print(f"  ✗ Failed to create Word document")
        
        # Convert to Excel
        excel_output = os.path.join(output_dir, f"{base_name}.xlsx")
        print("  Creating Excel spreadsheet...")
        if create_excel_spreadsheet(text_content, excel_output):
            print(f"  ✓ Created Excel spreadsheet")
        else:
            print(f"  ✗ Failed to create Excel spreadsheet")
        
        successful_conversions += 1
    
    print(f"\n🎉 Conversion process completed!")
    print(f"Successfully processed {successful_conversions} out of {len(pdf_files)} PDF files.")
    print(f"Converted files are located in: {output_dir}")
    return True

def main():
    """Main function"""
    print("Batch PDF to Word/Excel Converter")
    print("=" * 35)
    
    # Check and install required packages
    if not install_and_import_packages():
        print("Cannot proceed without required packages.")
        return False
    
    # Import packages after installation
    try:
        import PyPDF2
        import pdfplumber
        from docx import Document
        import openpyxl
        print("All packages imported successfully.")
    except ImportError as e:
        print(f"Failed to import packages: {e}")
        return False
    
    # Set directories
    input_dir = os.getcwd()  # Current directory
    output_dir = os.path.join(input_dir, "Converted_Files")
    
    print(f"Input directory: {input_dir}")
    print(f"Output directory: {output_dir}")
    print()
    
    try:
        process_pdf_files(input_dir, output_dir)
    except Exception as e:
        print(f"\nError during conversion process: {str(e)}")
        print("Traceback:")
        traceback.print_exc()
        return False
    
    return True

if __name__ == "__main__":
    success = main()
    if not success:
        sys.exit(1)